from docx import Document
import re
import json
import os
from datetime import datetime

DOC_TYPE_KEYWORDS = {
    "Articles of Association": ["articles of association", "articles of association (aoa)", "aoa", "articles"],
    "Memorandum of Association": ["memorandum of association", "memorandum", "moa"],
    "Register of Members and Directors": ["register of members", "register of directors", "register of members and directors"],
    "Incorporation Application Form": ["incorporation application", "application form"],
    "UBO Declaration Form": ["ubo declaration", "ultimate beneficial owner"],
}

def load_docx(path):
    doc = Document(path)
    paras = [p for p in doc.paragraphs]
    return doc, paras

def classify_document(paras):
    joined = " ".join([p.text for p in paras]).lower()
    for doc_type, keywords in DOC_TYPE_KEYWORDS.items():
        for kw in keywords:
            if kw in joined:
                return doc_type
    return "Unknown"

def find_issues(paras, doc_type):
    issues = []
    joined_text = " ".join([p.text for p in paras]).lower()

    if re.search(r"\b(uae federal|federal court|federal courts|dubai court|abu dhabi court)\b", joined_text, re.I):
        issues.append({
            "section": "Jurisdiction",
            "issue": "Document references non-ADGM courts (UAE Federal/Dubai/Abu Dhabi courts).",
            "severity": "High",
            "suggestion": "Replace jurisdiction reference with 'Abu Dhabi Global Market (ADGM)' where appropriate.",
            "citation": "ADGM Companies Regulations 2020",
            "para_index": None
        })

    tail = " ".join([p.text for p in paras[-8:]]).lower()
    if not re.search(r"\b(signature|signed by|for and on behalf|date:)\b", tail):
        issues.append({
            "section": "Signatory",
            "issue": "No obvious signature block or date found at end of document.",
            "severity": "High",
            "suggestion": "Add a signature block with printed name, signature, capacity and date.",
            "citation": "ADGM Registration checklist",
            "para_index": len(paras)-1
        })

    may_count = len(re.findall(r"\bmay\b", joined_text))
    shall_count = len(re.findall(r"\bshall\b", joined_text))
    if may_count > 3 and shall_count == 0:
        issues.append({
            "section": "Binding language",
            "issue": f"Document uses 'may' often ({may_count} times) and has no 'shall'—may be non-binding or ambiguous.",
            "severity": "Medium",
            "suggestion": "Use clearer mandatory language (e.g., 'shall') for obligations where intended.",
            "citation": "ADGM Companies Regulations 2020",
            "para_index": None
        })

    if doc_type == "Incorporation Application Form" or "incorporation" in joined_text:
        expected = {"articles of association", "memorandum of association", "register of members and directors",
                    "incorporation application form", "ubo declaration form"}
        present = set()
        for k in expected:
            if k in joined_text:
                present.add(k)
        missing = list(expected - present)
        if missing:
            issues.append({
                "section": "Checklist",
                "issue": f"Missing documents from incorporation checklist: {missing}",
                "severity": "High",
                "suggestion": "Upload the missing documents (e.g., Register of Members and Directors).",
                "citation": "ADGM Registration Checklist",
                "para_index": None
            })

    if "adgm" not in joined_text:
        issues.append({
            "section": "Jurisdiction",
            "issue": "Document never mentions 'ADGM' or 'Abu Dhabi Global Market'.",
            "severity": "Medium",
            "suggestion": "Confirm and add ADGM references where jurisdiction matters.",
            "citation": "ADGM Companies Regulations 2020",
            "para_index": None
        })

    return issues

def add_comments_and_save(original_path, issues, output_path):
    document = Document(original_path)
    for issue in issues:
        pi = issue.get("para_index")
        comment_text = f"{issue['issue']} Suggestion: {issue['suggestion']} (Ref: {issue['citation']})"
        if pi is not None and 0 <= pi < len(document.paragraphs):
            para = document.paragraphs[pi]
            try:
                document.add_comment(runs=para.runs, text=comment_text, author="CorporateAgent", initials="CA")
            except Exception:
                p = document.add_paragraph("[Review note]")
                try:
                    document.add_comment(runs=p.runs, text=comment_text, author="CorporateAgent", initials="CA")
                except Exception:
                    document.add_paragraph("NOTE: " + comment_text)
        else:
            p = document.add_paragraph("[Review note]")
            try:
                document.add_comment(runs=p.runs, text=comment_text, author="CorporateAgent", initials="CA")
            except Exception:
                document.add_paragraph("NOTE: " + comment_text)
    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    document.save(output_path)

def process_file(path, out_dir="reviewed_out"):
    doc, paras = load_docx(path)
    doc_type = classify_document(paras)
    issues = find_issues(paras, doc_type)
    base = os.path.splitext(os.path.basename(path))[0]
    ts = datetime.now().strftime("%Y%m%d%H%M%S")
    out_docx = os.path.join(out_dir, f"{base}_reviewed_{ts}.docx")
    os.makedirs(out_dir, exist_ok=True)
    add_comments_and_save(path, issues, out_docx)
    report = {
        "filename": os.path.basename(path),
        "process": "Company Incorporation" if "incorporation" in doc_type.lower() or "incorporation" in "".join([p.text for p in paras]).lower() else "Unknown",
        "document_type_detected": doc_type,
        "issues_found": issues,
        "reviewed_file": out_docx
    }
    return report
